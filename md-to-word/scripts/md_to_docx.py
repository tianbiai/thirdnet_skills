#!/usr/bin/env python3
"""
Markdown to Word (.docx) Converter

将 Markdown 文件转换为 Microsoft Word 文档。

Usage:
    python md_to_docx.py <input.md> [output.docx]
"""

import sys
import os
import re
from pathlib import Path
import io
import urllib.request
import urllib.error

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    import markdown
    from markdown.extensions import fenced_code, tables
    from PIL import Image
except ImportError as e:
    print(f"错误: 缺少必要的依赖。请运行: pip install python-docx markdown Pillow")
    print(f"详细信息: {e}")
    sys.exit(1)


class MarkdownToDocxConverter:
    """Markdown 转 Word 转换器"""

    def __init__(self, input_file_path: str = None):
        self.doc = Document()
        self._setup_styles()
        self.input_file_path = input_file_path  # 用于解析相对路径的图片
        self.input_dir = os.path.dirname(os.path.abspath(input_file_path)) if input_file_path else os.getcwd()

    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体为宋体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'  # 西文字体
        font.size = Pt(11)
        # 设置中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置标题样式也使用宋体
        for i in range(1, 4):
            try:
                heading_style = self.doc.styles[f'Heading {i}']
                heading_style.font.name = '宋体'
                heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            except KeyError:
                pass

    def _set_run_font(self, run, font_name='宋体', font_size=None):
        """
        为 run 设置统一的字体（同时处理中英文）

        Args:
            run: docx text run 对象
            font_name: 字体名称，默认宋体
            font_size: 字体大小（Pt），可选
        """
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = Pt(font_size)

    def _resolve_image_path(self, image_path: str) -> str:
        """
        解析图片路径，返回绝对路径

        Args:
            image_path: Markdown 中的图片路径（可以是相对路径、绝对路径或 URL）

        Returns:
            绝对路径，如果是 URL 则返回原值
        """
        # 如果是 URL，直接返回
        if image_path.startswith('http://') or image_path.startswith('https://'):
            return image_path

        # 如果是绝对路径，直接返回
        if os.path.isabs(image_path):
            return image_path

        # 相对路径：基于 Markdown 文件所在目录解析
        return os.path.join(self.input_dir, image_path)

    def _load_image(self, image_source: str) -> tuple:
        """
        加载图片并返回 PIL Image 对象和图片数据流

        Args:
            image_source: 图片路径或 URL

        Returns:
            (PIL Image 对象, BytesIO 流) 或 (None, None) 如果加载失败
        """
        try:
            # URL 图片
            if image_source.startswith('http://') or image_source.startswith('https://'):
                # 下载图片
                req = urllib.request.Request(
                    image_source,
                    headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
                )
                with urllib.request.urlopen(req, timeout=30) as response:
                    image_data = response.read()
                image_stream = io.BytesIO(image_data)
                img = Image.open(image_stream)
                return img, image_stream

            # 本地文件
            if os.path.exists(image_source):
                with open(image_source, 'rb') as f:
                    image_data = f.read()
                image_stream = io.BytesIO(image_data)
                img = Image.open(image_stream)
                return img, image_stream

            return None, None

        except Exception as e:
            print(f"  [警告] 无法加载图片 '{image_source}': {e}")
            return None, None

    def _add_image_to_doc(self, image_path: str, alt_text: str, max_width: float = 6.0) -> bool:
        """
        将图片添加到 Word 文档

        Args:
            image_path: 图片路径或 URL
            alt_text: 图片的替代文本
            max_width: 最大宽度（英寸），默认 6 英寸

        Returns:
            True 如果成功添加，False 如果失败
        """
        resolved_path = self._resolve_image_path(image_path)
        img, image_stream = self._load_image(resolved_path)

        if img is None:
            return False

        try:
            # 计算适当的尺寸，保持宽高比
            img_width, img_height = img.size

            # 转换为英寸（假设 96 DPI）
            width_inches = img_width / 96
            height_inches = img_height / 96

            # 如果宽度超过最大宽度，按比例缩放
            if width_inches > max_width:
                scale = max_width / width_inches
                width_inches = max_width
                height_inches = height_inches * scale

            # 限制最大高度为 8 英寸
            if height_inches > 8:
                scale = 8 / height_inches
                height_inches = 8
                width_inches = width_inches * scale

            # 添加图片到文档
            # 需要重置流位置
            image_stream.seek(0)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_stream, width=Inches(width_inches))

            # 如果有替代文本，添加图片说明
            if alt_text:
                caption = self.doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(alt_text)
                caption_run.italic = True
                self._set_run_font(caption_run, '宋体', 10)
                caption_run.font.color.rgb = RGBColor(100, 100, 100)

            return True

        except Exception as e:
            print(f"  [警告] 无法嵌入图片 '{image_path}': {e}")
            return False

    def convert(self, md_content: str, output_path: str):
        """
        转换 Markdown 内容为 Word 文档

        Args:
            md_content: Markdown 格式的文本内容
            output_path: 输出的 .docx 文件路径
        """
        # 解析 Markdown
        lines = md_content.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 处理代码块
            if stripped.startswith('```'):
                i = self._handle_code_block(lines, i)
                continue

            # 处理标题
            if stripped.startswith('#'):
                self._add_heading(stripped)
                i += 1
                continue

            # 忽略分隔线（章节间的分割线不转换到 Word）
            if stripped == '---' or stripped == '***' or stripped == '___':
                i += 1
                continue

            # 处理引用块
            if stripped.startswith('>'):
                i = self._handle_blockquote(lines, i)
                continue

            # 处理表格
            if self._is_table_row(stripped):
                i = self._handle_table(lines, i)
                continue

            # 处理列表
            if self._is_list_item(stripped):
                i = self._handle_list(lines, i)
                continue

            # 处理普通段落
            if stripped:
                self._add_paragraph(stripped)
                i += 1
            else:
                # 空行 - 作为段落分隔
                i += 1

        # 保存文档
        self.doc.save(output_path)
        print(f"[OK] 成功转换: {output_path}")

    def _is_table_row(self, line: str) -> bool:
        """检查是否为表格行（以 | 开头和结尾）"""
        stripped = line.strip()
        return stripped.startswith('|') and stripped.endswith('|')

    def _is_table_separator(self, line: str) -> bool:
        """检查是否为表格分隔行（包含 :--- | :---: | ---: 等）"""
        stripped = line.strip()
        if not self._is_table_row(stripped):
            return False
        # 移除首尾的 |，然后检查是否只包含 - : 和 |
        content = stripped[1:-1].strip()
        # 分隔行应该只包含 - : | 和空格
        return bool(re.match(r'^[\s|\-:]+$', content)) and '-' in content

    def _parse_table_row(self, line: str) -> list:
        """解析表格行，返回单元格内容列表"""
        stripped = line.strip()
        # 移除首尾的 |
        content = stripped[1:-1]
        # 按 | 分割
        cells = [cell.strip() for cell in content.split('|')]
        return cells

    def _handle_table(self, lines: list, start_idx: int) -> int:
        """处理表格"""
        i = start_idx
        table_rows = []

        # 收集所有表格行
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not self._is_table_row(stripped):
                break

            # 跳过分隔行（第二行通常是 :---|:---| 这种）
            if len(table_rows) == 1 and self._is_table_separator(stripped):
                i += 1
                continue

            cells = self._parse_table_row(stripped)
            table_rows.append(cells)
            i += 1

        if not table_rows:
            return i

        # 创建 Word 表格
        # 确定列数（取最大行宽度）
        num_cols = max(len(row) for row in table_rows) if table_rows else 1
        num_rows = len(table_rows)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        # 填充表格内容
        for row_idx, row_data in enumerate(table_rows):
            row = table.rows[row_idx]
            for col_idx in range(num_cols):
                cell = row.cells[col_idx]
                if col_idx < len(row_data):
                    content = self._format_inline(row_data[col_idx])
                else:
                    content = ''

                # 清空默认段落并使用 run 设置字体
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(content)
                self._set_run_font(run, '宋体')

                # 第一行作为表头，加粗显示
                if row_idx == 0:
                    run.bold = True

        # 在表格后添加空行
        self.doc.add_paragraph()

        return i

    def _is_list_item(self, line: str) -> bool:
        """检查是否为列表项"""
        # 无序列表: -, *, +
        # 有序列表: 1., 2), 等
        return bool(re.match(r'^[\s]*[-*+][\s]', line)) or \
               bool(re.match(r'^[\s]*\d+[.)][\s]', line))

    def _get_list_indent(self, line: str) -> int:
        """获取列表项的缩进级别"""
        stripped = line.lstrip()
        indent = len(line) - len(stripped)
        return indent // 2  # 每2个空格为一级

    def _handle_list(self, lines: list, start_idx: int) -> int:
        """处理列表（支持嵌套）"""
        i = start_idx

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 检查是否还是列表项或空行（在列表内部）
            if stripped == '':
                # 检查下一行是否还是列表
                if i + 1 < len(lines) and self._is_list_item(lines[i + 1]):
                    i += 1
                    continue
                else:
                    break

            if not self._is_list_item(line) and not line.startswith('    ') and not line.startswith('\t'):
                # 不是列表项且不是缩进内容
                break

            # 解析列表项
            indent = self._get_list_indent(line)
            content = self._extract_list_content(line)
            formatted_content = self._format_inline(content)

            # 添加列表项 - 使用 run 确保字体生效
            if re.match(r'^[\s]*\d+[.)]', line.lstrip()):
                # 有序列表
                p = self.doc.add_paragraph(style='List Number')
            else:
                # 无序列表
                p = self.doc.add_paragraph(style='List Bullet')

            # 添加带字体的 run
            run = p.add_run(formatted_content)
            self._set_run_font(run, '宋体')

            # 设置缩进级别
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))

            i += 1

        return i

    def _extract_list_content(self, line: str) -> str:
        """提取列表项的内容（去掉标记）"""
        # 去掉前导空白和列表标记
        stripped = line.lstrip()
        # 匹配无序列表标记
        match = re.match(r'^[-*+][\s](.*)', stripped)
        if match:
            return match.group(1)
        # 匹配有序列表标记
        match = re.match(r'^\d+[.)][\s](.*)', stripped)
        if match:
            return match.group(1)
        return stripped

    def _handle_blockquote(self, lines: list, start_idx: int) -> int:
        """处理引用块"""
        i = start_idx
        content_lines = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if stripped.startswith('>'):
                # 提取引用内容（去掉 > 符号）
                content = stripped[1:].strip()
                content_lines.append(content)
                i += 1
            elif stripped == '' and i + 1 < len(lines) and lines[i + 1].strip().startswith('>'):
                # 空行但在引用块内部
                content_lines.append('')
                i += 1
            else:
                break

        # 添加引用段落
        content = ' '.join(content_lines)
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(self._format_inline(content))
        run.italic = True
        self._set_run_font(run, '宋体')
        run.font.color.rgb = RGBColor(100, 100, 100)

        return i

    def _handle_code_block(self, lines: list, start_idx: int) -> int:
        """处理代码块"""
        i = start_idx + 1  # 跳过起始 ```
        code_lines = []

        # 提取语言标识（如果有）
        lang = lines[start_idx].strip()[3:].strip()

        while i < len(lines):
            line = lines[i]
            if line.strip() == '```':
                i += 1
                break
            code_lines.append(line)
            i += 1

        # 添加代码块
        code_content = '\n'.join(code_lines)
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(code_content)
        # 代码块也使用宋体，保持字体统一
        self._set_run_font(run, '宋体', 10)
        run.font.color.rgb = RGBColor(50, 50, 50)

        return i

    def _add_heading(self, line: str):
        """添加标题"""
        # 计算标题级别
        level = 0
        for char in line:
            if char == '#':
                level += 1
            else:
                break

        content = line[level:].strip()
        formatted_content = self._format_inline(content)

        # 添加标题 - 统一使用段落+run 方式以确保字体生效
        p = self.doc.add_paragraph()
        run = p.add_run(formatted_content)
        run.bold = True
        # 根据标题级别设置字号
        font_sizes = {1: 16, 2: 14, 3: 12, 4: 11, 5: 11, 6: 10}
        font_size = font_sizes.get(level, 11)
        self._set_run_font(run, '宋体', font_size)
        # 设置标题样式
        p.style = self.doc.styles[f'Heading {min(level, 3)}']

    def _add_paragraph(self, line: str):
        """添加普通段落，支持图片"""
        # 检查是否包含图片
        img_pattern = r'!\[(.*?)\]\((.*?)\)'

        # 查找所有图片
        images = list(re.finditer(img_pattern, line))

        if not images:
            # 没有图片，按普通段落处理 - 使用 run 确保字体生效
            formatted_content = self._format_inline(line)
            p = self.doc.add_paragraph()
            run = p.add_run(formatted_content)
            self._set_run_font(run, '宋体')
            return

        # 处理包含图片的段落
        # 策略：将图片和文本分开处理
        last_end = 0

        for match in images:
            # 添加图片前的文本
            text_before = line[last_end:match.start()]
            if text_before.strip():
                formatted_text = self._format_inline(text_before)
                p = self.doc.add_paragraph()
                run = p.add_run(formatted_text)
                self._set_run_font(run, '宋体')

            # 添加图片
            alt_text = match.group(1)
            image_path = match.group(2)

            success = self._add_image_to_doc(image_path, alt_text)
            if not success:
                # 图片加载失败，显示占位符
                p = self.doc.add_paragraph()
                run = p.add_run(f'[图片: {alt_text}]')
                run.italic = True
                self._set_run_font(run, '宋体')
                run.font.color.rgb = RGBColor(150, 150, 150)

            last_end = match.end()

        # 添加最后一段文本
        text_after = line[last_end:]
        if text_after.strip():
            formatted_text = self._format_inline(text_after)
            p = self.doc.add_paragraph()
            run = p.add_run(formatted_text)
            self._set_run_font(run, '宋体')

    def _format_inline(self, text: str) -> str:
        """
        处理行内格式标记，返回纯文本
        注意：python-docx 不支持完整的内联格式混合，这里进行简化处理
        """
        # 移除格式标记但保留内容
        # 粗体 **text** 或 __text__
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)

        # 斜体 *text*（只处理 * 包裹的情况，不处理 _ 包裹的情况）
        # 注意：不处理 _text_ 格式的斜体，因为在技术文档中下划线常用于命名（如 t_account_name）
        text = re.sub(r'\*(.*?)\*', r'\1', text)

        # 行内代码 `code`
        text = re.sub(r'`(.*?)`', r'\1', text)

        # 链接 [text](url) - 只保留文本
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)

        # 图片 ![alt](url) - 转换为 [图片: alt]
        text = re.sub(r'!\[(.*?)\]\(.*?\)', r'[图片: \1]', text)

        return text


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: python md_to_docx.py <input.md> [output.docx]")
        print("")
        print("参数:")
        print("  input.md    - 输入的 Markdown 文件路径")
        print("  output.docx - 输出的 Word 文件路径（可选，默认与输入同名）")
        sys.exit(1)

    input_path = sys.argv[1]

    # 确定输出路径
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        # 默认与输入同名，扩展名改为 .docx
        output_path = str(Path(input_path).with_suffix('.docx'))

    # 检查输入文件
    if not os.path.exists(input_path):
        print(f"错误: 输入文件不存在: {input_path}")
        sys.exit(1)

    # 读取 Markdown 内容
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except Exception as e:
        print(f"错误: 无法读取文件: {e}")
        sys.exit(1)

    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建目录: {output_dir}")

    # 转换
    try:
        converter = MarkdownToDocxConverter(input_path)
        converter.convert(md_content, output_path)
    except Exception as e:
        print(f"错误: 转换失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
